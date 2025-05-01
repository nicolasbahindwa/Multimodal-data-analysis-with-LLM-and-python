# src/readers/doc_reader.py
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
import tempfile
import subprocess

from ..utils.error_handler import ReadError, error_handler
from ..utils.logger import logger
from .base_reader import BaseReader


class DocReader(BaseReader):
    """
    Reader for Microsoft Word document files (.doc, .docx).
    Extracts content, structure, and metadata from Word documents.
    """
    
    @property
    def supported_extensions(self) -> list[str]:
        """List of supported file extensions."""
        return [".doc", ".docx"]
    
    @error_handler(
        error_type=ReadError,
        message="Failed to read Word document {args[0]}",
        log_traceback=True,
        raise_error=True
    )
    def _read_file(self, file_path: Path, **kwargs) -> Tuple[Any, Optional[Dict[str, Any]]]:
        """
        Read the content of a Word document.
        
        Args:
            file_path: Path to the Word document.
            **kwargs: Additional parameters, including:
                - extract_tables: Whether to extract tables (default is True).
                - extract_headers_footers: Whether to extract headers and footers (default is True).
                - extract_images: Whether to extract image references (default is False).
                - include_style_info: Whether to include style information (default is False).
                - preserve_formatting: Whether to preserve formatting like bold, italic (default is False).
                
        Returns:
            A tuple containing:
                - A dictionary with the document content and structure
                - A dictionary with metadata about the document
        """
        extract_tables = kwargs.get("extract_tables", True)
        extract_headers_footers = kwargs.get("extract_headers_footers", True)
        extract_images = kwargs.get("extract_images", False)
        include_style_info = kwargs.get("include_style_info", False)
        preserve_formatting = kwargs.get("preserve_formatting", False)
        
        ext = file_path.suffix.lower()
        
        if ext == ".docx":
            return self._read_docx(
                file_path, 
                extract_tables=extract_tables,
                extract_headers_footers=extract_headers_footers,
                extract_images=extract_images,
                include_style_info=include_style_info,
                preserve_formatting=preserve_formatting
            )
        elif ext == ".doc":
            return self._read_doc(file_path)
        else:
            raise ReadError(f"Unsupported file extension: {ext}")
    
    def _read_docx(
        self, 
        file_path: Path, 
        extract_tables: bool = True,
        extract_headers_footers: bool = True,
        extract_images: bool = False,
        include_style_info: bool = False,
        preserve_formatting: bool = False
    ) -> Tuple[Dict[str, Any], Dict[str, Any]]:
        """
        Read a .docx file using python-docx library.
        
        Args:
            file_path: Path to the .docx file.
            extract_tables: Whether to extract tables.
            extract_headers_footers: Whether to extract headers and footers.
            extract_images: Whether to extract image references.
            include_style_info: Whether to include style information.
            preserve_formatting: Whether to preserve formatting like bold, italic.
            
        Returns:
            A tuple containing:
                - A dictionary with the document content and structure
                - A dictionary with metadata about the document
        """
        try:
            import docx
        except ImportError:
            raise ReadError("python-docx package is required to read .docx files. Install it with 'pip install python-docx'.")
        
        logger.info(f"Reading .docx file: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if preserve_formatting:
                    # This is a simplified version, more complex formatting would need more work
                    text = ""
                    for run in para.runs:
                        if run.bold:
                            text += f"**{run.text}**"
                        elif run.italic:
                            text += f"*{run.text}*"
                        elif run.underline:
                            text += f"_{run.text}_"
                        else:
                            text += run.text
                    paragraphs.append(text)
                else:
                    paragraphs.append(para.text)
            
            # Extract tables if requested
            tables = []
            if extract_tables:
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    tables.append(table_data)
            
            # Extract headers and footers if requested
            headers = []
            footers = []
            if extract_headers_footers:
                for section in doc.sections:
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text:
                                headers.append(para.text)
                    
                    if section.footer:
                        for para in section.footer.paragraphs:
                            if para.text:
                                footers.append(para.text)
            
            # Extract image references if requested
            image_refs = []
            if extract_images:
                # This is a simplified approach - docx doesn't make it easy to extract images
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        image_refs.append(rel.target_ref)
            
            # Extract style info if requested
            styles = []
            if include_style_info:
                for style in doc.styles:
                    if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
                        styles.append(style.name)
            
            # Create combined content as plain text
            full_text = "\n".join(paragraphs)
            
            # Extract document properties/metadata
            doc_properties = {}
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                if hasattr(props, 'author') and props.author:
                    doc_properties["author"] = props.author
                if hasattr(props, 'title') and props.title:
                    doc_properties["title"] = props.title
                if hasattr(props, 'created') and props.created:
                    doc_properties["created"] = props.created
                if hasattr(props, 'modified') and props.modified:
                    doc_properties["modified"] = props.modified
                if hasattr(props, 'subject') and props.subject:
                    doc_properties["subject"] = props.subject
                if hasattr(props, 'keywords') and props.keywords:
                    doc_properties["keywords"] = props.keywords
                if hasattr(props, 'category') and props.category:
                    doc_properties["category"] = props.category
                if hasattr(props, 'comments') and props.comments:
                    doc_properties["comments"] = props.comments
            
            # Calculate word and character counts
            word_count = len(full_text.split())
            character_count = len(full_text)
            
            # Prepare content dictionary
            content = {
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables,
                "headers": headers,
                "footers": footers,
                "image_references": image_refs if extract_images else [],
                "styles": styles if include_style_info else []
            }
            
            # Prepare metadata
            metadata = {
                "word_count": word_count,
                "character_count": character_count,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                **doc_properties
            }
            
            logger.info(f"Successfully read .docx file {file_path}: {len(paragraphs)} paragraphs, {len(tables)} tables")
            
            return content, metadata
            
        except Exception as e:
            raise ReadError(f"Error reading .docx file {file_path}: {str(e)}", original_error=e)
    
    def _read_doc(self, file_path: Path) -> Tuple[Dict[str, Any], Dict[str, Any]]:
        """
        Read a .doc file using a combination of approaches (textract or antiword).
        
        Args:
            file_path: Path to the .doc file.
            
        Returns:
            A tuple containing:
                - A dictionary with the document content as text
                - A dictionary with metadata about the document
        """
        # Try multiple methods to extract text from .doc files
        text = None
        method_used = None
        
        # Method 1: Try using textract if available
        if text is None:
            try:
                import textract
                logger.info(f"Reading .doc file with textract: {file_path}")
                text = textract.process(str(file_path), extension="doc").decode("utf-8")
                method_used = "textract"
            except (ImportError, Exception) as e:
                logger.warning(f"Failed to read .doc with textract: {str(e)}")
        
        # Method 2: Try using mammoth if available
        if text is None:
            try:
                import mammoth
                logger.info(f"Reading .doc file with mammoth: {file_path}")
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
                method_used = "mammoth"
            except (ImportError, Exception) as e:
                logger.warning(f"Failed to read .doc with mammoth: {str(e)}")
        
        # Method 3: Try using antiword if available
        if text is None:
            try:
                logger.info(f"Reading .doc file with antiword: {file_path}")
                text = subprocess.check_output(["antiword", str(file_path)]).decode("utf-8")
                method_used = "antiword"
            except (subprocess.SubprocessError, FileNotFoundError) as e:
                logger.warning(f"Failed to read .doc with antiword: {str(e)}")
        
        # Method 4: Try converting to .docx and then reading it
        if text is None:
            try:
                # Check if libreoffice is available
                subprocess.check_output(["libreoffice", "--version"])
                
                with tempfile.TemporaryDirectory() as tmp_dir:
                    tmp_path = Path(tmp_dir)
                    output_file = tmp_path / f"{file_path.stem}.docx"
                    
                    logger.info(f"Converting .doc to .docx using LibreOffice: {file_path}")
                    subprocess.run([
                        "libreoffice", 
                        "--headless", 
                        "--convert-to", 
                        "docx", 
                        "--outdir", 
                        tmp_dir, 
                        str(file_path)
                    ], capture_output=True, check=True)
                    
                    if output_file.exists():
                        content, metadata = self._read_docx(output_file)
                        text = content["text"]
                        method_used = "libreoffice_conversion"
                    else:
                        logger.warning("LibreOffice conversion failed to produce output file")
            except (subprocess.SubprocessError, FileNotFoundError) as e:
                logger.warning(f"Failed to convert .doc with LibreOffice: {str(e)}")
        
        # If all methods failed, raise an error
        if text is None:
            raise ReadError(
                "Failed to read .doc file. Please install one of the following: "
                "textract, mammoth, antiword, or libreoffice."
            )
        
        # Estimate paragraphs by splitting on double newlines
        paragraphs = [p for p in text.split("\n\n") if p.strip()]
        
        # Calculate word and character counts
        word_count = len(text.split())
        character_count = len(text)
        
        # Create content dictionary
        content = {
            "text": text,
            "paragraphs": paragraphs,
            # Tables and other structured elements are not reliably extracted
            # from .doc files using these methods
            "tables": []
        }
        
        # Prepare metadata
        metadata = {
            "word_count": word_count,
            "character_count": character_count,
            "paragraph_count": len(paragraphs),
            "table_count": 0,  # We don't have reliable table info
            "content_type": "application/msword",
            "extraction_method": method_used
        }
        
        logger.info(f"Successfully read .doc file {file_path} using {method_used}: {len(paragraphs)} paragraphs")
        
        return content, metadata